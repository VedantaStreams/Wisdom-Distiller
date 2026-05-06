import os
import re
import math
import subprocess
import tempfile
from pathlib import Path

CHUNK_MB = 24


# ── Register Unicode fonts at module load time ────────────────────────────────
def _register_unicode_fonts():
    """Register Unicode-capable fonts. Called once at module load."""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    candidates = [
        ("UniSerif", "UniSerifBold",
         "/usr/share/fonts/truetype/freefont/FreeSerif.ttf",
         "/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf"),
        ("UniSans", "UniSansBold",
         "/usr/share/fonts/truetype/freefont/FreeSans.ttf",
         "/usr/share/fonts/truetype/freefont/FreeSansBold.ttf"),
        ("UniDejaVu", "UniDejaVuBold",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
    ]
    for reg_name, bold_name, reg_path, bold_path in candidates:
        if os.path.exists(reg_path):
            try:
                pdfmetrics.registerFont(TTFont(reg_name, reg_path))
                if os.path.exists(bold_path):
                    pdfmetrics.registerFont(TTFont(bold_name, bold_path))
                else:
                    bold_name = reg_name
                return reg_name, bold_name
            except Exception:
                continue
    return "Helvetica", "Helvetica-Bold"

# Register globally when module loads
try:
    _BASE_FONT, _BOLD_FONT = _register_unicode_fonts()
except Exception:
    _BASE_FONT, _BOLD_FONT = "Helvetica", "Helvetica-Bold"


# ── Audio duration ─────────────────────────────────────────────────────────────

def get_duration(path: str) -> float:
    """Get audio duration. Falls back gracefully if ffprobe not available."""
    # Method 1: ffprobe
    try:
        r = subprocess.run(
            ["ffprobe", "-v", "error", "-show_entries", "format=duration",
             "-of", "default=noprint_wrappers=1:nokey=1", path],
            capture_output=True, text=True, timeout=30
        )
        val = float(r.stdout.strip())
        if val > 0:
            return val
    except Exception:
        pass

    # Method 2: ffmpeg stderr
    try:
        r = subprocess.run(
            ["ffmpeg", "-i", path],
            capture_output=True, text=True, timeout=30
        )
        m = re.search(r"Duration:\s*(\d+):(\d+):([\d.]+)", r.stderr)
        if m:
            h, mn, s = int(m.group(1)), int(m.group(2)), float(m.group(3))
            return h * 3600 + mn * 60 + s
    except Exception:
        pass

    # Method 3: estimate from file size (128kbps assumption)
    return os.path.getsize(path) / 16000.0


# ── Audio splitting ────────────────────────────────────────────────────────────

def split_audio_ffmpeg(path: str) -> list:
    """Split audio into chunks under 25MB using ffmpeg.
    If file is already small enough, return as-is without calling ffmpeg."""
    file_size_mb = os.path.getsize(path) / (1024 * 1024)

    # Small file — no splitting needed, no ffmpeg call needed
    if file_size_mb <= CHUNK_MB:
        return [path]

    # Large file — need to split using ffmpeg
    duration = get_duration(path)
    n_chunks = math.ceil(file_size_mb / CHUNK_MB)
    chunk_sec = math.ceil(duration / n_chunks)
    tmp_dir = tempfile.mkdtemp()
    chunks = []

    for i in range(n_chunks):
        start = i * chunk_sec
        out = os.path.join(tmp_dir, f"chunk_{i+1}.mp3")
        result = subprocess.run(
            ["ffmpeg", "-y", "-i", path,
             "-ss", str(start), "-t", str(chunk_sec),
             "-acodec", "libmp3lame", "-q:a", "4", out],
            capture_output=True
        )
        if os.path.exists(out) and os.path.getsize(out) > 0:
            chunks.append(out)

    # If splitting failed for any reason, return original file
    if not chunks:
        return [path]

    return chunks


def extract_audio_from_video(video_path: str) -> str:
    """Extract audio from video file using ffmpeg."""
    tmp_dir = tempfile.mkdtemp()
    out = os.path.join(tmp_dir, "extracted_audio.mp3")
    subprocess.run(
        ["ffmpeg", "-y", "-i", video_path,
         "-vn", "-acodec", "libmp3lame", "-q:a", "4", out],
        capture_output=True
    )
    return out


def clean_youtube_url(url: str) -> str:
    """Strip playlist parameters, keep only video ID."""
    import urllib.parse
    parsed = urllib.parse.urlparse(url.strip())
    params = urllib.parse.parse_qs(parsed.query)
    video_id = params.get("v", [None])[0]
    if video_id:
        return f"https://www.youtube.com/watch?v={video_id}"
    return url.strip()


def prepare_audio_chunks(uploaded_files: list) -> list:
    """Save uploaded audio files to disk and split if needed."""
    all_chunks = []
    tmp_dir = tempfile.mkdtemp()
    for i, f in enumerate(uploaded_files):
        suffix = Path(f.name).suffix
        tmp_path = os.path.join(tmp_dir, f"file_{i+1}{suffix}")
        with open(tmp_path, "wb") as out:
            out.write(f.read())
        all_chunks.extend(split_audio_ffmpeg(tmp_path))
    return all_chunks


# ── Transcription ──────────────────────────────────────────────────────────────

def transcribe_chunks(chunks: list, openai_key: str,
                      progress_bar=None, status_text=None,
                      speaker: str = "", scripture: str = "") -> str:
    """Transcribe audio chunks via OpenAI Whisper API.
    Uses prompt hints to improve Sanskrit and spiritual term recognition."""
    from openai import OpenAI
    client = OpenAI(api_key=openai_key)
    parts = []

    # Build a prompt hint to help Whisper recognize spiritual/Sanskrit terms
    base_prompt = (
        "This is a Vedanta spiritual discourse. "
        "Common terms include: Brahman, Atman, Maya, Upanishad, Bhagavad Gita, "
        "Vedanta, Advaita, Guru, Shishya, Sanskrit, shloka, pranam, "
        "Paravidya, Aparavidya, jnana, karma, dharma, moksha, samsara, "
        "Chinmaya Mission, Swami, prana, manas, satyam, annam."
    )
    if speaker:
        base_prompt += f" The speaker is {speaker}."
    if scripture:
        base_prompt += f" The scripture being discussed is {scripture}."

    prev_text = ""
    for i, chunk_path in enumerate(chunks):
        if status_text:
            status_text.markdown(
                f"**Transcribing** part {i+1} of {len(chunks)}..."
            )
        if progress_bar:
            progress_bar.progress(0.05 + (i + 1) / len(chunks) * 0.60)

        # Use last 200 chars of previous chunk as context for continuity
        prompt = base_prompt
        if prev_text:
            prompt += " " + prev_text[-200:]

        with open(chunk_path, "rb") as audio_file:
            response = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file,
                response_format="text",
                prompt=prompt
            )
        text = response.strip()
        parts.append(text)
        prev_text = text

        try:
            if any(x in chunk_path for x in ["chunk_", "yt_audio", "extracted"]):
                os.unlink(chunk_path)
        except Exception:
            pass

    return "\n\n".join(parts)


# ── Summarization ──────────────────────────────────────────────────────────────

TRANSLITERATION_NOTE = (
    "When Sanskrit or regional language terms appear in the discourse, "
    "retain them and provide their English transliteration in parentheses "
    "immediately after. For example: Atman (the Self), Brahman (the Absolute), "
    "Maya (illusion). Keep all output in English. "
)

STYLE_PROMPTS = {
    "Bullet highlights": (
        TRANSLITERATION_NOTE
        + "Summarize the following discourse as clear bullet points grouped "
        + "under bold thematic headers. Each bullet should be concise."
    ),
    "Main takeaways": (
        TRANSLITERATION_NOTE
        + "List only the top 8-10 main takeaways from this discourse, "
        + "numbered, each in one or two sentences."
    ),
    "Detailed paragraphs": (
        TRANSLITERATION_NOTE
        + "Write a detailed, well-structured summary in prose paragraphs "
        + "covering all major themes, arguments, and conclusions."
    ),
    "Executive brief": (
        TRANSLITERATION_NOTE
        + "Write a crisp executive brief (max 300 words). Lead with the key "
        + "message, then main points and any action items."
    ),
    "Academic digest": (
        TRANSLITERATION_NOTE
        + "Produce an academic-style digest with sections: Overview, "
        + "Key Arguments, Evidence/Examples, Conclusions, and Notable Quotes."
    ),
    "Structured table": (
        "You will produce a structured table summary. "
        "Instructions are handled separately."
    ),
}

TABLE_COLUMNS = {
    "Main Point (verbatim)": "Extract the key point exactly as spoken.",
    "Explanation": "Explain what this point means in your own words.",
    "Example from discourse": "If an example was given, include it. Otherwise write N/A.",
    "Personal Reflection": "Leave this column empty for the reader to fill in.",
}


def summarize_text(transcript: str, style: str,
                   columns: list, anthropic_key: str) -> str:
    """Summarize transcript using Claude API."""
    import anthropic
    client = anthropic.Anthropic(api_key=anthropic_key)
    max_chars = 150_000

    def call_claude(prompt: str, text: str, max_tokens: int = 2000) -> str:
        if len(text) > max_chars:
            parts = [text[i:i+max_chars] for i in range(0, len(text), max_chars)]
            summaries = []
            for part in parts:
                msg = client.messages.create(
                    model="claude-sonnet-4-20250514",
                    max_tokens=1500,
                    messages=[{"role": "user",
                                "content": f"Summarize this section:\n\n{part}"}],
                )
                summaries.append(msg.content[0].text)
            combined = "\n\n".join(summaries)
            final = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=max_tokens,
                messages=[{"role": "user",
                            "content": f"{prompt}\n\n{combined}"}],
            )
            return final.content[0].text
        else:
            msg = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=max_tokens,
                messages=[{"role": "user",
                            "content": f"{prompt}\n\n{text}"}],
            )
            return msg.content[0].text

    if style == "Structured table":
        col_desc = "\n".join(
            [f"- {c}: {TABLE_COLUMNS[c]}"
             for c in columns if c in TABLE_COLUMNS]
        )
        prompt = (
            f"Create a markdown table from this discourse with these columns:\n"
            f"{col_desc}\n\n"
            f"Extract 10-15 key points. Format as a proper markdown table "
            f"with | separators. Leave the Personal Reflection column empty. "
            f"When Sanskrit or regional terms appear, add English meaning in "
            f"parentheses after them."
        )
        return call_claude(prompt, transcript, max_tokens=3000)
    else:
        prompt = STYLE_PROMPTS.get(style, STYLE_PROMPTS["Bullet highlights"])
        return call_claude(prompt, transcript)


# ── Translation ───────────────────────────────────────────────────────────────

LANGUAGES = {
    'English (default)': None,
    'Hindi (हिन्दी)': 'Hindi',
    'Kannada (ಕನ್ನಡ)': 'Kannada',
    'Telugu (తెలుగు)': 'Telugu',
    'Tamil (தமிழ்)': 'Tamil',
    'Marathi (मराठी)': 'Marathi',
    'Gujarati (ગુજરાતી)': 'Gujarati',
}

def translate_text(text, language, anthropic_key):
    """Translate text fully into target language. Sanskrit terms preserved as-is."""
    import anthropic
    client = anthropic.Anthropic(api_key=anthropic_key)
    max_chars = 150_000

    instruction = (
        "You are a translator. Translate ALL of the following text completely into "
        + language + ". "
        "Every word, label, heading, bullet point, and sentence must be in "
        + language + " only. "
        "Do NOT leave any English words in the output except for: "
        "Sanskrit terms (Atman, Brahman, Maya, Upanishads, Bhagavad Gita, etc.), "
        "proper nouns (names of people and places), "
        "and technical terms that have no equivalent in " + language + ". "
        "Preserve the structure, formatting, bullet points, and headings exactly. "
        "The entire output must read as natural, fluent " + language + "."
    )

    def call(t):
        msg = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": instruction + "\n\n" + t}],
        )
        return msg.content[0].text

    if len(text) > max_chars:
        parts = [text[i:i+max_chars] for i in range(0, len(text), max_chars)]
        return "\n\n".join(call(p) for p in parts)
    return call(text)



# ── Discourse Insights ────────────────────────────────────────────────────────

def analyze_discourse(
    transcript: str,
    anthropic_key: str,
    speaker_hint: str = "",
    topic_hint: str = "",
    scripture_hint: str = ""
) -> dict:
    """Analyze transcript for speaker, topic, scriptures, key terms.
    Optional hints provided by the user take priority over auto-detection."""
    import anthropic
    import json
    client = anthropic.Anthropic(api_key=anthropic_key)

    # Build context from user hints
    hints = ""
    if speaker_hint:
        hints += f"The speaker is: {speaker_hint}.\n"
    if topic_hint:
        hints += f"The topic/text is: {topic_hint}.\n"
    if scripture_hint:
        hints += f"The scripture being discussed is: {scripture_hint}.\n"
    if hints:
        hints = "Context provided by user:\n" + hints + "\n"

    prompt = (
        "Analyze the following spiritual discourse transcript and extract structured insights. "
        + hints +
        "Use the user-provided context above as authoritative — do not contradict it. "
        "For anything not provided, infer from the transcript. "
        "Respond ONLY in this exact JSON format with no extra text:\n"
        "{\n"
        '  "speaker": "Name of speaker",\n'
        '  "topic": "Main topic or theme in one sentence",\n'
        '  "scripture_text": "Primary scripture being discussed",\n'
        '  "scriptures": ["Specific verses referenced e.g. Bhagavad Gita 2.20"],\n'
        '  "key_terms": ["Key Sanskrit terms used"],\n'
        '  "tradition": "e.g. Advaita Vedanta / Chinmaya Mission / etc."\n'
        "}\n\n"
        "Transcript:\n" + transcript[:8000]
    )

    try:
        msg = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=600,
            messages=[{"role": "user", "content": prompt}]
        )
        text = msg.content[0].text.strip()
        text = text.replace("```json", "").replace("```", "").strip()
        return json.loads(text)
    except Exception:
        return {
            "speaker": speaker_hint or "Unknown",
            "topic": topic_hint or "Could not determine",
            "scripture_text": scripture_hint or "Unknown",
            "scriptures": [],
            "key_terms": [],
            "tradition": "Unknown"
        }


# ── PDF export ─────────────────────────────────────────────────────────────────

def _parse_markdown_table(content: str):
    """Parse markdown table into (headers, rows). Returns None if not a table."""
    lines = [l.strip() for l in content.strip().split("\n") if l.strip()]
    table_lines = [l for l in lines if l.startswith("|")]
    if len(table_lines) < 2:
        return None
    def is_sep(row):
        return all(set(c.replace("-","").replace(":","").strip()) <= {""} for c in row)
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if not is_sep(cells):
            rows.append(cells)
    if len(rows) < 2:
        return None
    return rows[0], rows[1:]


def make_pdf(title: str, content: str,
             speaker: str = "", topic: str = "",
             scripture: str = "", language: str = "",
             verses: list = None, key_terms: list = None) -> bytes:
    """Generate a PDF. Includes discourse header if details provided."""
    verses = verses or []
    key_terms = key_terms or []
    try:
        import weasyprint  # noqa
        return _make_pdf_weasyprint(title, content,
                                    speaker=speaker, topic=topic,
                                    scripture=scripture, language=language,
                                    verses=verses, key_terms=key_terms)
    except ImportError:
        return _make_pdf_reportlab(title, content,
                                   speaker=speaker, topic=topic,
                                   scripture=scripture, language=language,
                                   verses=verses, key_terms=key_terms)
    except Exception:
        return _make_pdf_reportlab(title, content,
                                   speaker=speaker, topic=topic,
                                   scripture=scripture, language=language,
                                   verses=verses, key_terms=key_terms)


def _weasyprint_write_pdf(html_string):
    """Write PDF with Telugu font subsetting bug workaround."""
    from weasyprint import HTML
    import fontTools.ttLib.tables.O_S_2f_2 as os2_module

    def safe_set_unicode_ranges(self, bits):
        bits = {b for b in bits if 0 <= b <= 122}
        original_set(self, bits)

    original_set = os2_module.table_O_S_2f_2.setUnicodeRanges
    os2_module.table_O_S_2f_2.setUnicodeRanges = safe_set_unicode_ranges
    try:
        return HTML(string=html_string).write_pdf()
    finally:
        os2_module.table_O_S_2f_2.setUnicodeRanges = original_set


def _make_pdf_weasyprint(title: str, content: str,
                         speaker: str = "", topic: str = "",
                         scripture: str = "", language: str = "",
                         verses: list = None, key_terms: list = None) -> bytes:
    """PDF via weasyprint — supports all Indian scripts natively."""
    from weasyprint import HTML
    verses = verses or []
    key_terms = key_terms or []

    parsed = _parse_markdown_table(content)
    is_table = parsed is not None

    if is_table:
        headers, rows = parsed
        header_cells = "".join(f"<th>{h}</th>" for h in headers)
        body_rows = ""
        for i, row in enumerate(rows):
            while len(row) < len(headers):
                row.append("")
            bg = "#f9f6f0" if i % 2 == 0 else "#ffffff"
            cells = "".join(f"<td>{c}</td>" for c in row)
            body_rows += f'<tr style="background:{bg};">{cells}</tr>'
        main_content = (
            f"<table><thead><tr>{header_cells}</tr></thead>"
            f"<tbody>{body_rows}</tbody></table>"
        )
        landscape = "size: A4 landscape;"
    else:
        paragraphs = ""
        for line in content.split("\n"):
            if line.startswith("## "):
                paragraphs += f"<h2>{line[3:]}</h2>"
            elif line.startswith("# "):
                paragraphs += f"<h1>{line[2:]}</h1>"
            elif line.startswith("- ") or line.startswith("* "):
                paragraphs += f"<li>{line[2:]}</li>"
            elif line.strip():
                paragraphs += f"<p>{line}</p>"
        main_content = paragraphs
        landscape = "size: A4 portrait;"

    import os as _os
    _noto_devanagari = '/usr/share/fonts/truetype/noto/NotoSerifDevanagari-Regular.ttf'
    _noto_kannada = '/usr/share/fonts/truetype/noto/NotoSansKannada-Regular.ttf'
    _noto_telugu = '/usr/share/fonts/truetype/noto/NotoSansTelugu-Regular.ttf'
    _noto_tamil = '/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf'
    _noto_gujarati = '/usr/share/fonts/truetype/noto/NotoSansGujarati-Regular.ttf'

    _font_faces = ""
    for _fname, _fpath in [
        ("NotoSerif", _noto_devanagari),
        ("NotoKannada", _noto_kannada),
        ("NotoTelugu", _noto_telugu),
        ("NotoTamil", _noto_tamil),
        ("NotoGujarati", _noto_gujarati),
    ]:
        if _os.path.exists(_fpath):
            _font_faces += f"@font-face {{ font-family: '{_fname}'; src: url('{_fpath}'); }}"

    # Build discourse header block for PDF
    header_html = ""
    if any([speaker, topic, scripture, verses, key_terms]):
        def _cell(label, value, color="#1a1a1a"):
            return (
                f"<div style='flex:1;padding:0.6rem 1rem;'>"
                f"<div style='font-size:7pt;color:#888;text-transform:uppercase;"
                f"letter-spacing:0.8px;margin-bottom:3px;'>{label}</div>"
                f"<div style='font-size:11pt;font-weight:700;color:{color};'>"
                f"{value or chr(8212)}</div></div>"
            )
        verse_str = " · ".join(verses) if verses else "—"
        terms_str = " · ".join(key_terms) if key_terms else "—"
        header_html = (
            "<div style='border-top:4px solid #c9a96e;border-radius:8px;"
            "background:#f9f6f0;margin-bottom:16px;'>"
            "<div style='font-size:7pt;color:#c9a96e;text-transform:uppercase;"
            "letter-spacing:1px;font-weight:600;padding:8px 16px 4px;'>"
            "✦ Discourse Details</div>"
            "<div style='display:flex;border-top:1px solid #e8ddd0;'>"
            + _cell("🎙️ Speaker", speaker)
            + _cell("📖 Topic", topic)
            + _cell("📚 Scripture", scripture, "#8B6914")
            + _cell("🌐 Language", language or "English (default)")
            + "</div>"
            + (
                "<div style='border-top:1px solid #e8ddd0;padding:6px 16px;'>"
                f"<span style='font-size:7pt;color:#888;text-transform:uppercase;"
                f"letter-spacing:0.8px;'>📜 Verses Referenced</span>&nbsp;"
                f"<span style='font-size:9pt;color:#8B6914;'>{verse_str}</span>"
                f"<br/>"
                f"<span style='font-size:7pt;color:#888;text-transform:uppercase;"
                f"letter-spacing:0.8px;'>🔑 Key Terms</span>&nbsp;"
                f"<span style='font-size:9pt;color:#555;font-style:italic;'>{terms_str}</span>"
                "</div>"
                if (verses or key_terms) else ""
            )
            + "</div></div>"
        )

    html = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>
{_font_faces}
@page {{ {landscape} margin: 1.5cm; }}
body {{
    font-family: 'NotoSerif', 'NotoKannada', 'NotoTelugu', 'NotoTamil',
                 'NotoGujarati', FreeSerif, FreeSans, serif;
    font-size: 10px;
    color: #111;
    line-height: 1.6;
}}
h1 {{ font-size: 18px; color: #c9a96e; text-align: center; margin-bottom: 16px; }}
h2 {{ font-size: 14px; color: #333; margin-top: 12px; }}
p {{ margin: 4px 0 8px; }}
li {{ margin: 3px 0; }}
table {{
    width: 100%;
    border-collapse: collapse;
    font-size: 9px;
    page-break-inside: auto;
}}
thead tr {{
    background: #c9a96e;
    color: white;
}}
th {{
    padding: 7px 8px;
    text-align: left;
    font-weight: 600;
    font-size: 9px;
    border: 1px solid #b8935a;
}}
td {{
    padding: 6px 8px;
    border: 1px solid #ddd;
    vertical-align: top;
    line-height: 1.5;
}}
tr {{ page-break-inside: avoid; }}
</style>
</head>
<body>
<h1>{title}</h1>
{header_html}
{main_content}
</body>
</html>"""

    return _weasyprint_write_pdf(html)


def _make_pdf_reportlab(title: str, content: str,
                        speaker: str = "", topic: str = "",
                        scripture: str = "", language: str = "",
                        verses: list = None, key_terms: list = None) -> bytes:
    """Fallback PDF via reportlab with FreeSerif font."""
    verses = verses or []
    key_terms = key_terms or []
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle)
    from reportlab.lib.enums import TA_CENTER
    import io

    base_font, bold_font = _BASE_FONT, _BOLD_FONT
    parsed = _parse_markdown_table(content)
    is_table = parsed is not None
    pagesize = landscape(A4) if is_table else A4
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=pagesize,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("T", parent=styles["Title"], fontName=bold_font,
                                  fontSize=16, textColor=colors.HexColor("#c9a96e"),
                                  spaceAfter=12, alignment=TA_CENTER)
    body_style = ParagraphStyle("B", parent=styles["Normal"], fontName=base_font,
                                 fontSize=10, leading=16, spaceAfter=8,
                                 textColor=colors.HexColor("#222222"))
    cell_style = ParagraphStyle("C", parent=styles["Normal"], fontName=base_font,
                                 fontSize=8.5, leading=13, wordWrap="CJK")
    header_style = ParagraphStyle("H", parent=styles["Normal"], fontName=bold_font,
                                   fontSize=9, leading=12,
                                   textColor=colors.white, alignment=TA_CENTER)
    story = [Paragraph(title, title_style), Spacer(1, 0.4*cm)]

    # Add discourse header if details provided
    if any([speaker, topic, scripture, verses, key_terms]):
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors as rl_colors
        verse_str = " · ".join(verses) if verses else "—"
        terms_str = " · ".join(key_terms) if key_terms else "—"
        page_w = (landscape(A4)[0] if is_table else A4[0]) - 3*cm
        hdr_data = [[
            Paragraph(f"<b>🎙️ Speaker</b><br/>{speaker or '—'}", body_style),
            Paragraph(f"<b>📖 Topic</b><br/>{topic or '—'}", body_style),
            Paragraph(f"<b>📚 Scripture</b><br/>{scripture or '—'}", body_style),
            Paragraph(f"<b>🌐 Language</b><br/>{language or 'English'}", body_style),
        ]]
        htbl = Table(hdr_data, colWidths=[page_w/4]*4)
        htbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,-1), rl_colors.HexColor("#f9f6f0")),
            ("LINEABOVE",  (0,0), (-1,0),  3, rl_colors.HexColor("#c9a96e")),
            ("GRID",       (0,0), (-1,-1), 0.5, rl_colors.HexColor("#e8ddd0")),
            ("TOPPADDING",    (0,0), (-1,-1), 8),
            ("BOTTOMPADDING", (0,0), (-1,-1), 8),
            ("LEFTPADDING",   (0,0), (-1,-1), 10),
        ]))
        story.append(htbl)
        if verses or key_terms:
            vtbl = Table([[
                Paragraph(f"<b>📜 Verses Referenced:</b> {verse_str}", body_style),
                Paragraph(f"<b>🔑 Key Terms:</b> <i>{terms_str}</i>", body_style),
            ]], colWidths=[page_w/2, page_w/2])
            vtbl.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,-1), rl_colors.HexColor("#f9f6f0")),
                ("GRID",       (0,0), (-1,-1), 0.5, rl_colors.HexColor("#e8ddd0")),
                ("TOPPADDING",    (0,0), (-1,-1), 6),
                ("BOTTOMPADDING", (0,0), (-1,-1), 6),
                ("LEFTPADDING",   (0,0), (-1,-1), 10),
            ]))
            story.append(vtbl)
        story.append(Spacer(1, 0.4*cm))
    if is_table:
        headers, rows = parsed
        table_data = [[Paragraph(h, header_style) for h in headers]]
        for row in rows:
            while len(row) < len(headers):
                row.append("")
            table_data.append([Paragraph(c, cell_style) for c in row])
        page_w = landscape(A4)[0] - 3*cm
        col_w = page_w / len(headers)
        tbl = Table(table_data, colWidths=[col_w]*len(headers), repeatRows=1)
        tbl.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#c9a96e")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("VALIGN",     (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1),
             [colors.HexColor("#f9f6f0"), colors.HexColor("#ffffff")]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#cccccc")),
            ("TOPPADDING",    (0,0), (-1,-1), 5),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING",   (0,0), (-1,-1), 6),
            ("RIGHTPADDING",  (0,0), (-1,-1), 6),
        ]))
        story.append(tbl)
    else:
        for line in content.split("\n"):
            if line.strip():
                safe = (line.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;"))
                story.append(Paragraph(safe, body_style))
            else:
                story.append(Spacer(1, 0.2*cm))
    doc.build(story)
    return buffer.getvalue()


def make_docx(title: str, content: str) -> bytes:
    """Generate a DOCX from text content. Renders markdown tables as real Word tables."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import io

    doc = Document()

    # Set landscape for tables
    parsed = _parse_markdown_table(content)
    is_table = parsed is not None

    if is_table:
        section = doc.sections[0]
        section.orientation = 1  # landscape
        section.page_width, section.page_height = section.page_height, section.page_width

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0xC9, 0xA9, 0x6E)
        run.font.size = Pt(18)
        run.font.name = "Arial Unicode MS"  # Unicode font for heading

    doc.add_paragraph()

    if is_table:
        headers, rows = parsed

        # Add Word table
        num_cols = len(headers)
        tbl = doc.add_table(rows=1 + len(rows), cols=num_cols)
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            run = hdr_cells[i].paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Arial Unicode MS"
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Gold background for header
            tc = hdr_cells[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "C9A96E")
            tcPr.append(shd)

        # Data rows
        for r_idx, row in enumerate(rows):
            while len(row) < num_cols:
                row.append("")
            row_cells = tbl.rows[r_idx + 1].cells
            # Alternating row background
            fill_color = "F9F6F0" if r_idx % 2 == 0 else "FFFFFF"
            for c_idx, cell_text in enumerate(row):
                row_cells[c_idx].text = cell_text
                run = row_cells[c_idx].paragraphs[0].runs
                if run:
                    run[0].font.size = Pt(8)
                    run[0].font.name = "Arial Unicode MS"  # Supports Sanskrit
                # Row background
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), fill_color)
                tcPr.append(shd)

        # Auto-fit columns
        for col in tbl.columns:
            for cell in col.cells:
                cell.width = Cm(5)

    else:
        for line in content.split("\n"):
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(line[2:])
            elif line.strip():
                doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Markdown table to HTML ─────────────────────────────────────────────────────

def markdown_table_to_html(md: str) -> str:
    """Convert a markdown table into a styled Excel-like HTML table."""
    lines = [l.strip() for l in md.strip().split("\n") if l.strip()]
    table_lines = [l for l in lines if l.startswith("|")]

    if not table_lines:
        return "<div class=\"output-box\">" + md + "</div>"

    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        rows.append(cells)

    # Remove separator rows like |---|---|
    def is_separator(row):
        return all(set(c.replace("-", "").replace(":", "").strip()) <= {""} for c in row)

    rows = [r for r in rows if not is_separator(r)]

    if not rows:
        return "<div class=\"output-box\">" + md + "</div>"

    header = rows[0]
    body = rows[1:]

    th_cells = "".join(f"<th>{h}</th>" for h in header)
    tbody_rows = ""
    for i, row in enumerate(body):
        while len(row) < len(header):
            row.append("")
        bg = "#161616" if i % 2 == 0 else "#1a1a1a"
        td_cells = "".join(f"<td>{c}</td>" for c in row)
        tbody_rows += f"<tr style=\"background:{bg};\">{td_cells}</tr>"

    return (
        "<div style=\"overflow-x:auto; margin-top:1rem;\">"
        "<table style=\"width:100%; border-collapse:collapse; font-size:0.85rem;"
        " font-family:'DM Sans',sans-serif; color:#d4c9b8;"
        " border:1px solid #2a2a2a; border-radius:8px; overflow:hidden;\">"
        "<thead>"
        "<tr style=\"background:#1e1a14; border-bottom:2px solid #c9a96e;\">"
        + th_cells
        + "</tr></thead><tbody>"
        + tbody_rows
        + "</tbody></table></div>"
    )


TABLE_CSS = (
    "<style>"
    "table th { padding: 10px 14px; text-align: left; color: #c9a96e;"
    " font-weight: 500; font-size: 0.82rem; letter-spacing: 0.4px;"
    " text-transform: uppercase; border-right: 1px solid #2a2a2a; }"
    "table td { padding: 9px 14px; border-right: 1px solid #1e1e1e;"
    " border-bottom: 1px solid #1e1e1e; vertical-align: top; line-height: 1.6; }"
    "table tr:hover td { background: #222 !important; }"
    "table td:last-child, table th:last-child { border-right: none; }"
    "</style>"
)

